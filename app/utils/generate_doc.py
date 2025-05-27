from docx import Document
import tempfile
from app.utils.templates import build_cv_html

# To convert plain cv text to a styled docx file
def create_docx_file(name:str, email:str, content:str) -> str:
    doc = Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(email)
    doc.add_paragraph("")

    # To split by section and format
    for section in content.split("\n\n"):
        if section.startswith("**"):
            section = section.strip("*").strip()
            doc.add_heading(section, level=1)
        else:
            doc.add_paragraph(section)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx")as tmp_file:
        doc.save(tmp_file.name)
        return tmp_file.name


