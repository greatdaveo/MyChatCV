from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from utils.templates import build_cv_html

# To convert plain cv text to a styled docx file
def create_docx_file(name: str, email: str, content: str, phone: str = "", location: str = "", linkedin: str = "") -> str:
    doc = Document()
    # To set up the docx margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # To add header with name
    header = doc.add_heading(name, level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_info = []
    if email:
        contact_info.append(email)
    if phone:
        contact_info.append(phone)
    if location:
        contact_info.append(location)
    if linkedin:
        contact_info.append(linkedin)

    if contact_info:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.add_run(" • ".join(contact_info))
        contact_para.runs[0].font.size = Pt(10)

    # To add spacing
    doc.add_paragraph()

    # To process content sections
    sections = content.split('\n\n')

    for section in sections:
        section = section.strip()
        if not section:
            continue

        # To check if this is a section header,(starts with *)
        if section.startswith('**') and section.endswith('**'):
            section_title = section.strip('*').strip()
            heading = doc.add_heading(section_title, level=1)
            heading.style.font.size = Pt(14)
            heading.style.font.bold = True
        else:
            # To process regular content
            process_docx_content(doc, section)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        doc.save(tmp_file.name)
        return tmp_file.name


# Process content for DOCX formatting with better structure.
def process_docx_content(doc, content: str):
    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # To check for bullet points or achievements
        if line.startswith('•') or line.startswith('-'):
            # Create a bullet point
            para = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            para.runs[0].font.size = Pt(10)
        else:
            # For normal paragraph
            para = doc.add_paragraph(line)
            para.runs[0].font.size = Pt(10)

    # To add spacing after section
    doc.add_paragraph()